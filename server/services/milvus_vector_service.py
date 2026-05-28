"""
文档向量化服务（Milvus）
负责文档解析、文本分块与 Milvus 向量存储；接口与 Chroma 版 vector_service 保持一致。
"""
import os
import time
from abc import ABC
from typing import Any, List, Optional

from flask import current_app
from huggingface_hub.cli import _output
from langchain_community.retrievers import BM25Retriever
from langchain_core.callbacks import CallbackManagerForRetrieverRun
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pydantic import ConfigDict
from pymilvus import MilvusClient
from sentence_transformers import SentenceTransformer

from services import vector_service
from utils.ensemble_retriever import EnsembleRetriever

os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")


class MilvusDenseRetriever(BaseRetriever, ABC):
    """LangChain检索器：Milvus 稠密向量相似度检索。"""

    def __init__(self, *args: Any, **kwargs: Any):
        super().__init__(*args, **kwargs)
        model_config = ConfigDict(arbitrary_types_allowed=True)
        self.vector_service = Any
        self.kb_id = int
        self.k = 10

    def _get_relevant_documents(
            self, query: str, *, run_manager: CallbackManagerForRetrieverRun
    ) -> list[Document]:
        return self.vector_service.search_dense(self.kb_id, query, k=self.k)


def _output_fields():
    fields = ["text", "doc_id", "file_name", "chunk_index"]
    window_key = current_app.config.get("SENTENCE_WINDOW_METADATA_KEY", "window")
    if window_key not in fields:
        fields.append(window_key)

    return fields


def _row_to_document(entry):
    text = entry.get("text") or ""
    meta = {}
    for key, value in entry.items():
        if key in ("text", "vector", "id", "distance"):
            continue
        if value is not None:
            meta[key] = value
    return Document(page_content=text, metadata=meta)


def _get_collection_name(kb_id):
    return f"kb_{kb_id}"


def _load_file(file_path, file_type):
    text = ""
    if file_type in ("txt", "md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    elif file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    elif file_type == "docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

    return text


class MilvusVectorService:
    def __init__(self):
        """初始化嵌入模型，文本分块器，Milvus 客户端及检索相关缓存"""
        model_path = current_app.config["EMBEDDING_MODEL_PATH"]
        self.embeddings = SentenceTransformer(model_path)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=current_app.config["CHUNK_SIZE"],
            chunk_overlap=current_app.config["CHUNK_OVERLAP"],
            length_function=len)
        self.client = MilvusClient(uri=current_app.config["MILVUS_URI"])
        self.embedding_dim = int(current_app.config.get("EMBEDDING_DIM"))
        self.batch_size = current_app.config.get("BATCH_SIZE", 10)
        self.max_retries = current_app.config.get("EMBED_MAX_RETRIES", 3)
        self.nlist = int(current_app.config.get("MILVUS_INDEX_NLIST", 128))
        self.nprobe = int(current_app.config.get("MILVUS_INDEX_NPROBE", 32))
        self._query_batch_size = int(current_app.config.get("MILVUS_QUERY_BATCH_SIZE", 500))
        self._loaded_collections: set[str] = set()
        self._kb_documents_cache = {}
        self._kb_bm25_retriever_cache = {}

    def get_retriever(self, kb_id: int):
        return MilvusDenseRetriever(vector_service=self, kb_id=kb_id, k=current_app.config["RETRIEVER_TOP_K"])

    def search_dense(self, kb_id, query, k) -> List[Document]:
        collection_name = self._ensure_collection(kb_id)
        stats = self.client.get_collection_stats(collection_name)
        if int(stats.get("row_count", 0)) == 0:
            return []
        query_vec = self._encode_query(query)
        results = self.client.search(collection_name=collection_name,
                                     data=[query_vec],
                                     limit=k,
                                     output_fields=_output_fields(),
                                     search_params={"metric_type": "IP", "params": {"nprobe": self.nprobe}},
                                     anns_fields="vector")
        docs = []
        for hits in results:
            for hit in hits:
                entry = hit.get("entity") or {}
                doc = _row_to_document(entry)
                distance = hit.get("distance")
                if distance is not None:
                    doc.metadata["relevance_score"] = float(distance)
                docs.append(doc)

        return docs

    def _ensure_collection(self, kb_id):
        collection_name = _get_collection_name(kb_id)
        if not self.client.has_collection(collection_name):
            index_params = self.client.prepare_index_params()
            index_params.add_index(field_name="vector",
                                   index_type="FLAT",
                                   metric_type="IP",
                                   params={"nlist": self.nlist})
            self.client.create_collection(collection_name=collection_name,
                                          dimension=self.embedding_dim,
                                          metric_type="IP",
                                          auto_id=True,
                                          enable_dynamic_field=True,
                                          index_params=index_params
                                          )
            current_app.logger.info("已创建 Milvus 集合：%s", collection_name)
        if collection_name not in self._loaded_collections:
            self.client.load_collection(collection_name)
            self._loaded_collections.add(collection_name)

        return collection_name

    def _encode_query(self, query):
        prefixed = f'{query}'
        vec = self.embeddings.encode([prefixed], normalize_embeddings=True, show_progress_bar=False)

        return vec[0].tolist()

    def _encode_documents(self, texts):
        vectors = self.embeddings.encode(texts, normalize_embeddings=True, show_progress_bar=False)
        return vectors.tolist()

    def process_document(self, doc_id, file_path, file_type, kb_id, display_file_name):
        text = _load_file(file_path, file_type)
        if not text.strip():
            raise ValueError("文档内容为空，无法进行向量化")

        file_name = display_file_name or os.path.basename(file_path)

        chunks = self.text_splitter.split_text(text)
        if not chunks:
            raise ValueError("文档分块失败")
        metadatas = [
            {
                "doc_id": doc_id,
                "file_name": file_name,
                "chunk_index": i
            }
            for i in range(len(chunks))
        ]
        collection_name = self._ensure_collection(kb_id)
        for i in range(0, len(chunks), self.batch_size):
            batch_end = min(i + self.batch_size, len(chunks))
            batch_chunks = chunks[i:batch_end]
            batch_meta = metadatas[i:batch_end]
            vectors = self._encode_documents(batch_chunks)
            rows = []
            for vec, chunk_text, meta in zip(vectors, batch_chunks, batch_meta):
                row = {"vector": vec, "text": chunk_text, **meta}
                rows.append(row)
            self._insert_with_retry(collection_name, rows)
        self._invalidate_kb_documents_cache(kb_id)
        self.client.flush(collection_name)
        return len(chunks)

    def _insert_with_retry(self, collection_name, rows):
        last_error = None
        for attempt in range(self.max_retries):
            try:
                self.client.insert(collection_name, rows)
                return
            except Exception as e:
                last_error = e
                if attempt == self.max_retries - 1:
                    raise
                wait = 2 ** attempt
                current_app.logger.warning("Milvus 插入失败(第%s次)，%s秒后重试: %s", attempt + 1, wait, e)
                time.sleep(wait)
            raise last_error

    def _invalidate_kb_documents_cache(self, kb_id):
        self._kb_documents_cache.pop(kb_id, None)
        self._kb_bm25_retriever_cache.pop(kb_id, None)

    def delete_document(self, doc_id, kb_id):
        collection_name = _get_collection_name(kb_id)
        if not self.client.has_collection(collection_name):
            return
        self.client.delete(collection_name=collection_name,
                           filter=f"doc_id=={int(doc_id)}")
        self.client.flush(collection_name)
        self._invalidate_kb_documents_cache(kb_id)
