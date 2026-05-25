"""
文档向量化服务（Milvus）
负责文档解析、文本分块与 Milvus 向量存储；接口与 Chroma 版 vector_service 保持一致。
"""
import os
import time
from typing import Any, List, Optional

from flask import current_app
from langchain_community.retrievers import BM25Retriever
from langchain_core.callbacks import CallbackManagerForRetrieverRun
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pydantic import ConfigDict
from pymilvus import MilvusClient
from sentence_transformers import SentenceTransformer

from utils.ensemble_retriever import EnsembleRetriever

os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")

BGE_QUERY_PREFIX = "为这个句子生成表示以用于检索相关文章："


def _get_collection_name(kb_id: int) -> str:
    return f"kb_{kb_id}"


def _load_file(file_path: str, file_type: str) -> str:
    text = ""
    if file_type in ("txt", "md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    elif file_type == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    elif file_type == "docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
    return text


def _chinese_sentence_splitter(text: str) -> List[str]:
    import re

    if not text or not text.strip():
        return []
    parts = re.split(r"(?<=[。！？!?；;\n])", text)
    sentences = [p.strip() for p in parts if p and p.strip()]
    if not sentences:
        return [text.strip()]
    return sentences


def _row_to_document(self, row: dict) -> Document:
    text = row.get("text") or ""
    meta = {}
    for key, value in row.items():
        if key in ("text", "vector", "id", "distance"):
            continue
        if value is not None:
            meta[key] = value
    return Document(page_content=text, metadata=meta)


def _output_fields() -> List[str]:
    fields = ["text", "doc_id", "file_name", "chunk_index"]
    window_key = current_app.config.get("SENTENCE_WINDOW_METADATA_KEY", "window")
    if window_key not in fields:
        fields.append(window_key)
    return fields


def _sentence_window_chunks(self, text: str):
    """使用 SentenceWindow 策略将全文切分为句子块及对应上下文窗口。"""
    from llama_index.core import Document as LIDocument
    from llama_index.core.node_parser import SentenceWindowNodeParser

    window_size = int(current_app.config.get("SENTENCE_WINDOW_SIZE", 3))
    window_key = current_app.config.get("SENTENCE_WINDOW_METADATA_KEY", "window")
    use_zh = current_app.config.get("SENTENCE_WINDOW_CHINESE_SPLIT", True)

    parser_kwargs = {
        "window_size": max(1, window_size),
        "window_metadata_key": window_key,
        "include_prev_next_rel": False,
    }
    if use_zh:
        parser_kwargs["sentence_splitter"] = _chinese_sentence_splitter

    parser = SentenceWindowNodeParser.from_defaults(**parser_kwargs)
    li_doc = LIDocument(text=text)
    nodes = parser.get_nodes_from_documents([li_doc])

    chunks, windows = [], []
    for node in nodes:
        sentence = (getattr(node, "text", None) or "").strip()
        if not sentence:
            continue
        window_text = node.metadata.get(window_key) or sentence
        chunks.append(sentence)
        windows.append(window_text)

    return chunks, windows, window_key


class MilvusDenseRetriever(BaseRetriever):
    """LangChain 检索器：Milvus 稠密向量相似度检索。"""

    model_config = ConfigDict(arbitrary_types_allowed=True)

    vector_service: Any
    kb_id: int
    k: int = 10

    def _get_relevant_documents(
            self, query: str, *, run_manager: CallbackManagerForRetrieverRun
    ) -> List[Document]:
        return self.vector_service.search_dense(self.kb_id, query, k=self.k)


class MilvusVectorService:
    def __init__(self):
        """初始化嵌入模型、文本分块器、Milvus 客户端及检索相关缓存。"""
        model_path = current_app.config["EMBEDDING_MODEL_PATH"]
        self.embeddings = SentenceTransformer(model_path)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=current_app.config["CHUNK_SIZE"],
            chunk_overlap=current_app.config["CHUNK_OVERLAP"],
            length_function=len,
        )
        self.client = MilvusClient(uri=current_app.config["MILVUS_URI"])
        self.embedding_dim = int(current_app.config["EMBEDDING_DIM"])
        self.batch_size = current_app.config.get("EMBED_BATCH_SIZE", 10)
        self.max_retries = current_app.config.get("EMBED_MAX_RETRIES", 3)
        self.nlist = int(current_app.config.get("MILVUS_INDEX_NLIST", 128))
        self.nprobe = int(current_app.config.get("MILVUS_SEARCH_NPROBE", 32))
        self._query_batch_size = int(current_app.config.get("MILVUS_QUERY_BATCH_SIZE", 500))
        self._loaded_collections: set[str] = set()
        self._kb_documents_cache: dict = {}
        self._kb_bm25_retriever_cache: dict = {}

    def get_retriever(self, kb_id: int):
        """按配置返回稠密向量检索器或稠密+BM25 混合检索器。"""
        if current_app.config.get("HYBRID_RETRIEVAL"):
            return self._get_hybrid_retriever(kb_id)
        return MilvusDenseRetriever(
            vector_service=self,
            kb_id=kb_id,
            k=current_app.config["RETRIEVER_TOP_K"],
        )

    def _get_hybrid_retriever(self, kb_id: int):
        """组装 Milvus 稠密检索与 BM25 稀疏检索的加权融合检索器。"""
        dense_k = current_app.config.get("HYBRID_DENSE_K", 10)
        sparse_k = current_app.config.get("HYBRID_SPARSE_K", 10)
        weights = current_app.config.get("HYBRID_ENSEMBLE_WEIGHTS", [0.7, 0.3])

        dense = MilvusDenseRetriever(
            vector_service=self, kb_id=kb_id, k=dense_k
        )

        # 将知识库加载至内存
        sparse_docs = self._get_cached_kb_documents(kb_id)
        if not sparse_docs:
            return dense

        bm25 = self._kb_bm25_retriever_cache.get(kb_id)
        if bm25 is None:
            bm25 = BM25Retriever.from_documents(sparse_docs)
            self._kb_bm25_retriever_cache[kb_id] = bm25
        # 返回 top bm25.k (= HYBRID_SPARSE_K) 条
        bm25.k = sparse_k

        return EnsembleRetriever(
            retrievers=[dense, bm25],
            weights=list(weights),
        )

    def _get_cached_kb_documents(self, kb_id: int) -> List[Document]:
        """从内存缓存获取知识库全文块，未命中时从 Milvus 加载。"""
        if kb_id not in self._kb_documents_cache:
            self._kb_documents_cache[kb_id] = self._load_kb_documents_from_milvus(kb_id)

        return self._kb_documents_cache[kb_id]

    def _load_kb_documents_from_milvus(self, kb_id: int) -> List[Document]:
        """遍历 Milvus 集合，将全部文本块加载为 LangChain Document 列表。"""
        collection_name = _get_collection_name(kb_id)
        if not self.client.has_collection(collection_name):
            return []
        self._ensure_collection(kb_id)
        docs: List[Document] = []
        output_fields = _output_fields()

        try:
            iterator = self.client.query_iterator(
                collection_name=collection_name,
                batch_size=self._query_batch_size,
                output_fields=output_fields,
            )
            while True:
                batch = iterator.next()
                if not batch:
                    break
                for row in batch:
                    if row.get("text"):
                        docs.append(_row_to_document(row))
        except Exception:
            batch = self.client.query(
                collection_name=collection_name,
                filter="doc_id >= 0",
                output_fields=output_fields,
                limit=16384,
            )
            for row in batch:
                if row.get("text"):
                    docs.append(_row_to_document(row))
        return docs

    def search_dense(self, kb_id: int, query: str, k: int) -> List[Document]:
        """对查询文本做向量编码，在 Milvus 中按相似度返回 top-k 文档块。"""
        collection_name = self._ensure_collection(kb_id)
        stats = self.client.get_collection_stats(collection_name)
        if int(stats.get("row_count", 0)) == 0:
            return []

        query_vec = self._encode_query(query)
        results = self.client.search(
            collection_name=collection_name,
            data=[query_vec],
            limit=k,
            output_fields=_output_fields(),
            search_params={"metric_type": "IP", "params": {"nprobe": self.nprobe}},
            anns_field="vector",
        )
        docs: List[Document] = []
        for hits in results:
            for hit in hits:
                entity = hit.get("entity") or {}
                doc = _row_to_document(entity)
                distance = hit.get("distance")
                if distance is not None:
                    doc.metadata["relevance_score"] = float(distance)
                docs.append(doc)

        return docs

    def _ensure_collection(self, kb_id: int) -> str:
        """确保知识库对应 Milvus 集合已创建、建索引并加载到内存。"""
        collection_name = _get_collection_name(kb_id)
        if not self.client.has_collection(collection_name):
            index_params = self.client.prepare_index_params()
            index_params.add_index(
                field_name="vector",
                index_type="FLAT",
                metric_type="IP",
                params={"nlist": self.nlist},
            )
            self.client.create_collection(
                collection_name=collection_name,
                dimension=self.embedding_dim,
                metric_type="IP",
                auto_id=True,
                enable_dynamic_field=True,
                index_params=index_params,
            )
            current_app.logger.info("已创建 Milvus 集合: %s", collection_name)

        if collection_name not in self._loaded_collections:
            self.client.load_collection(collection_name)
            self._loaded_collections.add(collection_name)

        return collection_name

    def process_document(self, doc_id: int, file_path: str, file_type: str, kb_id: int, display_file_name: Optional[str] = None, ) -> int:
        """解析文件、分块、向量化并写入 Milvus，返回生成的块数量。"""
        text = _load_file(file_path, file_type)
        if not text.strip():
            raise ValueError("文档内容为空，无法进行向量化")

        file_name = display_file_name or os.path.basename(file_path)
        window_key = current_app.config.get("SENTENCE_WINDOW_METADATA_KEY", "window")

        if current_app.config.get("USE_SENTENCE_WINDOW"):
            chunks, windows, window_key = self._sentence_window_chunks(text)
            if not chunks:
                raise ValueError("文档分块失败（SentenceWindow）")
            metadatas = [
                {
                    "doc_id": doc_id,
                    "file_name": file_name,
                    "chunk_index": i,
                    window_key: windows[i],
                }
                for i in range(len(chunks))
            ]
        else:
            chunks = self.text_splitter.split_text(text)
            if not chunks:
                raise ValueError("文档分块失败")
            metadatas = [
                {"doc_id": doc_id, "file_name": file_name, "chunk_index": i}
                for i in range(len(chunks))
            ]

        collection_name = self._ensure_collection(kb_id)

        for i in range(0, len(chunks), self.batch_size):
            batch_end = min(i + self.batch_size, len(chunks))
            batch_chunks = chunks[i:batch_end]
            batch_meta = metadatas[i:batch_end]
            vectors = self._encode_documents(batch_chunks)
            rows = []
            for vec, chunk_text, meta in zip(vectors, batch_chunks, batch_meta):
                row = {"vector": vec, "text": chunk_text, **meta}
                rows.append(row)
            self._insert_with_retry(collection_name, rows)

        self._invalidate_kb_documents_cache(kb_id)
        self.client.flush(collection_name)
        return len(chunks)

    def _insert_with_retry(self, collection_name: str, rows: List[dict]) -> None:
        """向 Milvus 批量插入向量行，失败时按指数退避重试。"""
        last_error = None
        for attempt in range(self.max_retries):
            try:
                self.client.insert(collection_name, rows)
                return
            except Exception as e:
                last_error = e
                if attempt == self.max_retries - 1:
                    raise
                wait = 2 ** attempt
                current_app.logger.warning(
                    "Milvus 插入失败(第%s次)，%s秒后重试: %s",
                    attempt + 1,
                    wait,
                    e,
                )
                time.sleep(wait)
        raise last_error

    def _encode_query(self, query: str) -> List[float]:
        """将用户查询编码为归一化稠密向量，供相似度检索使用。"""
        prefixed = f"{query}"
        vec = self.embeddings.encode(
            [prefixed], normalize_embeddings=True, show_progress_bar=False
        )
        return vec[0].tolist()

    def _encode_documents(self, texts: List[str]) -> List[List[float]]:
        """将一批文本块批量编码为归一化向量列表。"""
        vectors = self.embeddings.encode(
            texts, normalize_embeddings=True, show_progress_bar=False
        )
        return vectors.tolist()

    def _invalidate_kb_documents_cache(self, kb_id: int) -> None:
        """清除指定知识库的文档列表与 BM25 检索器内存缓存。"""
        self._kb_documents_cache.pop(kb_id, None)
        self._kb_bm25_retriever_cache.pop(kb_id, None)

    def delete_document(self, doc_id: int, kb_id: int) -> None:
        """按 doc_id 删除 Milvus 中该文档的全部向量块并刷新缓存。"""
        collection_name = _get_collection_name(kb_id)
        if not self.client.has_collection(collection_name):
            return
        self.client.delete(
            collection_name=collection_name,
            filter=f"doc_id == {int(doc_id)}",
        )
        self.client.flush(collection_name)
        self._invalidate_kb_documents_cache(kb_id)
