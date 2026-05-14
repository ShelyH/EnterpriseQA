"""
文档向量化服务
负责文档解析、文本分块和Chroma向量存储
"""
import os
import time
from flask import current_app
from langchain_chroma import Chroma
from utils.ensemble_retriever import EnsembleRetriever
from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

os.environ["HF_HUB_OFFLINE"] = "1"
os.environ["TRANSFORMERS_OFFLINE"] = "1"


def _get_collection_name(kb_id):
    """
    根据知识库ID生成Chroma集合名称
    每个知识库使用独立的collection进行隔离
    """
    return f"kb_{kb_id}"


def _load_file(file_path, file_type):
    """
    根据文件类型加载文档内容
    :param file_path: 文件路径
    :param file_type: 文件类型（txt/pdf/md/docx）
    :return: 文本内容
    """
    text = ''
    if file_type in ('txt', 'md'):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
    elif file_type == 'pdf':
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    elif file_type == 'docx':
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + '\n'

    return text


def _chinese_sentence_splitter(text: str):
    """
    按常见中文/英文句末标点切句，供 SentenceWindowNodeParser 使用。
    若未切出多句，则整段作为一句，避免空结果。
    """
    import re

    if not text or not text.strip():
        return []
    parts = re.split(r'(?<=[。！？!?；;\n])', text)
    sentences = [p.strip() for p in parts if p and p.strip()]
    if not sentences:
        return [text.strip()]
    return sentences


class VectorService:
    """文档向量化服务类"""

    def __init__(self):
        """初始化嵌入模型和文本分割器"""
        self.embeddings = HuggingFaceEmbeddings(
            model_name="bge-small-zh-v1.5",
            model_kwargs={"device": "cpu", "local_files_only": True},
            encode_kwargs={"normalize_embeddings": True},
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=current_app.config['CHUNK_SIZE'],
            chunk_overlap=current_app.config['CHUNK_OVERLAP'],
            length_function=len
        )
        self.persist_dir = current_app.config['CHROMA_PERSIST_DIR']
        self.batch_size = current_app.config.get('EMBED_BATCH_SIZE', 10)
        self.max_retries = current_app.config.get('EMBED_MAX_RETRIES', 3)
        self._vectorstores = {}  # 缓存已创建的 Chroma 实例
        # 混合检索：chunk 文本列表与 BM25 检索器缓存（写入/删除后失效）
        self._kb_documents_cache = {}
        self._kb_bm25_retriever_cache = {}

    def _sentence_window_chunks(self, text):
        """
        使用 LlamaIndex SentenceWindowNodeParser：每节点正文为一句（用于向量），
        metadata[window] 为前后若干句拼接（用于 RAG 上下文扩展）。
        """
        from llama_index.core import Document as LIDocument
        from llama_index.core.node_parser import SentenceWindowNodeParser

        window_size = int(current_app.config.get('SENTENCE_WINDOW_SIZE', 3))
        window_key = current_app.config.get('SENTENCE_WINDOW_METADATA_KEY', 'window')
        use_zh = current_app.config.get('SENTENCE_WINDOW_CHINESE_SPLIT', True)

        parser_kwargs = {
            'window_size': max(1, window_size),
            'window_metadata_key': window_key,
            'include_prev_next_rel': False,
        }
        if use_zh:
            parser_kwargs['sentence_splitter'] = _chinese_sentence_splitter

        parser = SentenceWindowNodeParser.from_defaults(**parser_kwargs)
        li_doc = LIDocument(text=text)
        nodes = parser.get_nodes_from_documents([li_doc])

        chunks = []
        windows = []
        for node in nodes:
            sentence = (getattr(node, 'text', None) or '').strip()
            if not sentence:
                continue
            window_text = node.metadata.get(window_key) or sentence
            chunks.append(sentence)
            windows.append(window_text)

        return chunks, windows, window_key

    def _add_texts_with_retry(self, vectorstore, texts, metadatas, ids):
        """
        :param vectorstore: Chroma向量库实例
        :param texts: 文本分块列表
        :param metadatas: 元数据列表
        :param ids: ID列表
        """
        last_error = None
        for attempt in range(self.max_retries):
            try:
                vectorstore.add_texts(texts=texts, metadatas=metadatas, ids=ids)
                return
            except Exception as e:
                last_error = e
                err_msg = str(e)
                is_retryable = any(code in err_msg for code in ('502', '503', '504'))
                if not is_retryable or attempt == self.max_retries - 1:
                    raise
                wait = 2 ** attempt
                current_app.logger.warning(
                    f'Ollama嵌入请求失败(第{attempt + 1}次)，{wait}秒后重试: {err_msg}'
                )
                time.sleep(wait)
        raise last_error

    def _invalidate_kb_documents_cache(self, kb_id):
        self._kb_documents_cache.pop(kb_id, None)
        self._kb_bm25_retriever_cache.pop(kb_id, None)

    def _get_vectorstore(self, kb_id):
        collection_name = _get_collection_name(kb_id)
        if kb_id not in self._vectorstores:
            self._vectorstores[kb_id] = Chroma(
                collection_name=collection_name,
                embedding_function=self.embeddings,
                persist_directory=self.persist_dir
            )
        return self._vectorstores[kb_id]

    def _load_kb_documents_from_chroma(self, kb_id):
        """从 Chroma 读出某知识库全部 chunk，供 BM25 使用。"""
        vectorstore = self._get_vectorstore(kb_id)
        coll = vectorstore._collection
        res = coll.get(include=['documents', 'metadatas'])
        texts = res.get('documents') or []
        metas = res.get('metadatas') or []
        docs = []
        for text, meta in zip(texts, metas):
            if text:
                docs.append(Document(page_content=text, metadata=dict(meta) if meta else {}))
        return docs

    def _get_cached_kb_documents(self, kb_id):
        if kb_id not in self._kb_documents_cache:
            self._kb_documents_cache[kb_id] = self._load_kb_documents_from_chroma(kb_id)
        return self._kb_documents_cache[kb_id]

    def _get_hybrid_retriever(self, kb_id):
        dense_k = current_app.config.get('HYBRID_DENSE_K', 10)
        sparse_k = current_app.config.get('HYBRID_SPARSE_K', 10)
        weights = current_app.config.get('HYBRID_ENSEMBLE_WEIGHTS', [0.6, 0.4])

        vectorstore = self._get_vectorstore(kb_id)
        dense = vectorstore.as_retriever(
            search_type='similarity',
            search_kwargs={'k': dense_k},
        )

        sparse_docs = self._get_cached_kb_documents(kb_id)
        if not sparse_docs:
            return dense

        bm25 = self._kb_bm25_retriever_cache.get(kb_id)
        if bm25 is None:
            bm25 = BM25Retriever.from_documents(sparse_docs)
            self._kb_bm25_retriever_cache[kb_id] = bm25
        bm25.k = sparse_k

        return EnsembleRetriever(
            retrievers=[dense, bm25],
            weights=list(weights),
        )

    def process_document(self, doc_id, file_path, file_type, kb_id, display_file_name=None):
        """
        处理文档：预检查 -> 解析文件 -> 文本分块 -> 分批存入向量库。
        USE_SENTENCE_WINDOW 为真时使用 LlamaIndex SentenceWindowNodeParser（按句向量、邻句窗口写入元数据）。
        :param doc_id: 文档ID
        :param file_path: 文件路径
        :param file_type: 文件类型
        :param kb_id: 知识库ID
        :param display_file_name 为写入向量元数据的展示名
        :return: 分块数量
        """

        text = _load_file(file_path, file_type)
        if not text.strip():
            raise ValueError('文档内容为空，无法进行向量化')

        file_name = display_file_name or os.path.basename(file_path)

        if current_app.config.get('USE_SENTENCE_WINDOW'):
            chunks, windows, window_key = self._sentence_window_chunks(text)
            if not chunks:
                raise ValueError('文档分块失败（SentenceWindow）')
            metadatas = [
                {
                    'doc_id': doc_id,
                    'file_name': file_name,
                    'chunk_index': i,
                    window_key: windows[i],
                }
                for i in range(len(chunks))
            ]
        else:
            chunks = self.text_splitter.split_text(text)
            if not chunks:
                raise ValueError('文档分块失败')
            metadatas = [{'doc_id': doc_id, 'file_name': file_name, 'chunk_index': i} for i in range(len(chunks))]

        ids = [f"doc_{doc_id}_chunk_{i}" for i in range(len(chunks))]

        vectorstore = self._get_vectorstore(kb_id)

        for i in range(0, len(chunks), self.batch_size):
            batch_end = min(i + self.batch_size, len(chunks))
            self._add_texts_with_retry(
                vectorstore,
                texts=chunks[i:batch_end],
                metadatas=metadatas[i:batch_end],
                ids=ids[i:batch_end],
            )

        self._invalidate_kb_documents_cache(kb_id)
        return len(chunks)

    def delete_document(self, doc_id, kb_id):
        """
        从向量库中删除指定文档的所有分块
        :param doc_id: 文档ID
        :param kb_id: 知识库ID
        """
        vectorstore = self._get_vectorstore(kb_id)
        # 根据文档ID过滤并删除
        vectorstore._collection.delete(where={'doc_id': doc_id})
        self._invalidate_kb_documents_cache(kb_id)

    def get_retriever(self, kb_id):
        """
        获取指定知识库的检索器（开启混合检索时为向量+BM25 融合）
        :param kb_id: 知识库ID
        :return: LangChain BaseRetriever
        """
        if current_app.config.get('HYBRID_RETRIEVAL'):
            return self._get_hybrid_retriever(kb_id)

        vectorstore = self._get_vectorstore(kb_id)
        return vectorstore.as_retriever(
            search_type='similarity',
            search_kwargs={
                'k': current_app.config['RETRIEVER_TOP_K']
            }
        )
